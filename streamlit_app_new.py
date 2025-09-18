import streamlit as st
import requests
import json
import time
from typing import List, Dict, Any
import pandas as pd
from io import BytesIO
import uuid

# Configure Streamlit page
st.set_page_config(
    page_title="Credit Review System",
    page_icon="üìä",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Backend API base URL
API_BASE_URL = "http://localhost:8000/api/v1"

def check_backend_health():
    """Check if backend is running"""
    try:
        response = requests.get(f"{API_BASE_URL}/health/", timeout=2)
        return response.status_code == 200
    except Exception as e:
        print(f"Backend health check failed: {e}")
        return False

def upload_document(file, document_name=None):
    """Upload document to backend"""
    files = {"file": file}
    data = {}
    if document_name:
        data["document_name"] = document_name
    
    try:
        response = requests.post(f"{API_BASE_URL}/documents/upload", files=files, data=data)
        return response.json() if response.status_code == 200 else None
    except Exception as e:
        st.error(f"Upload failed: {str(e)}")
        return None

def delete_document(document_id):
    """Delete document from backend"""
    try:
        response = requests.delete(f"{API_BASE_URL}/documents/{document_id}")
        return response.json() if response.status_code == 200 else None
    except Exception as e:
        st.error(f"Delete failed: {str(e)}")
        return None

def get_documents():
    """Get list of uploaded documents"""
    try:
        response = requests.get(f"{API_BASE_URL}/documents/")
        return response.json().get("documents", []) if response.status_code == 200 else []
    except:
        return []

def get_document_chunks(document_id, limit=2):
    """Get chunks for a specific document"""
    try:
        response = requests.get(f"{API_BASE_URL}/documents/{document_id}/chunks", params={"limit": limit})
        return response.json() if response.status_code == 200 else None
    except Exception as e:
        st.error(f"Failed to get document chunks: {str(e)}")
        return None

def fetch_public_company_documents(ticker_symbol, exchange, quarter=None, year=None, filing_types=None):
    """Fetch public company documents from SEC EDGAR or international equivalents"""
    try:
        data = {
            "ticker_symbol": ticker_symbol,
            "exchange": exchange,
            "quarter": quarter,
            "year": year,
            "filing_types": filing_types or []
        }
        response = requests.post(f"{API_BASE_URL}/documents/fetch-public", json=data)
        return response.json() if response.status_code == 200 else None
    except Exception as e:
        st.error(f"Failed to fetch public company documents: {str(e)}")
        return None

def upload_document_with_sonnet_fallback(file, document_name=None, processing_mode="comprehensive", focus_areas=None):
    """Upload document using Sonnet AI fallback processing"""
    files = {"file": file}
    data = {}
    if document_name:
        data["document_name"] = document_name
    if processing_mode:
        data["processing_mode"] = processing_mode
    if focus_areas:
        data["focus_areas"] = focus_areas

    try:
        response = requests.post(f"{API_BASE_URL}/documents/upload-sonnet-fallback", files=files, data=data)
        return response.json() if response.status_code == 200 else None
    except Exception as e:
        st.error(f"Sonnet fallback upload failed: {str(e)}")
        return None

def generate_segment_content(segment_data, selected_document_ids=None):
    """Generate content for a segment using the report coordinator"""
    try:
        data = {
            "segment_data": segment_data,
            "validation_enabled": True,
            "selected_document_ids": selected_document_ids or []
        }
        # Use the existing endpoint but with actual segment data
        response = requests.post(f"{API_BASE_URL}/segments/generate-report", json=data)
        return response.json() if response.status_code == 200 else None
    except Exception as e:
        return {"error": str(e)}

def create_report_and_export_word(report_data, include_validation_comments=True):
    """Create a report and export as Word document"""
    try:
        # First create a report record
        report_response = requests.post(f"{API_BASE_URL}/reports/", json={
            "title": report_data['title'],
            "description": f"Generated report with {len(report_data['sections'])} sections"
        })
        
        if report_response.status_code != 200:
            return {"error": "Failed to create report record"}
        
        report_id = report_response.json()['id']
        
        # Create segments for the report
        segment_ids = []
        validation_results_list = []
        
        for i, section in enumerate(report_data['sections']):
            # Get validation results for this section
            validation_results = section.get('validation_results', {})
            validation_results_list.append(validation_results)
            
            segment_data = {
                "report_id": report_id,
                "name": section['name'],
                "description": f"Section {i+1}",
                "prompt": f"Generated content for {section['name']}",
                "order_index": i,
                "required_document_types": [],
                "generation_settings": {},
                "generated_content": section['content'],
                "content_status": "completed",
                "validation_results": validation_results
            }
            
            segment_response = requests.post(f"{API_BASE_URL}/segments/", json=segment_data)
            if segment_response.status_code == 200:
                segment_ids.append(segment_response.json()['id'])
        
        # Export as Word document
        export_response = requests.post(
            f"{API_BASE_URL}/reports/{report_id}/export/word",
            params={"include_validations": include_validation_comments}
        )
        
        if export_response.status_code == 200:
            return export_response.json()
        else:
            return {"error": f"Export failed: {export_response.text}"}
            
    except Exception as e:
        return {"error": str(e)}

def download_word_document(download_url):
    """Download Word document from the backend"""
    try:
        # Construct full URL - handle both absolute and relative URLs
        if download_url.startswith('http'):
            full_url = download_url
        else:
            # Remove /api/v1 from base URL and add the download_url
            base_url = API_BASE_URL.replace('/api/v1', '')
            full_url = f"{base_url}{download_url}"

        response = requests.get(full_url, timeout=30)
        if response.status_code == 200:
            return response.content
        else:
            st.error(f"Download failed with status {response.status_code}: {response.text}")
            return None
    except Exception as e:
        st.error(f"Download failed: {str(e)}")
        return None

def create_basic_word_content(report):
    """Create a basic Word document as fallback when backend fails"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create new document
        doc = Document()

        # Add title
        title = doc.add_heading(report['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add generation date
        date_para = doc.add_paragraph(f"Generated on {report['generation_date']}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page break
        doc.add_page_break()

        # Add sections
        for i, section in enumerate(report['sections']):
            # Section heading
            doc.add_heading(f"{i+1}. {section['name']}", level=1)

            # Section content
            content = section.get('content', 'No content available')
            paragraphs = content.split('\n\n')

            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

            # Add spacing
            doc.add_paragraph()

        # Save to BytesIO
        from io import BytesIO
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)

        return word_buffer.getvalue()

    except ImportError:
        # If python-docx is not available, create a simple text-based "Word" document
        content = f"{report['title']}\n\nGenerated on {report['generation_date']}\n\n"

        for i, section in enumerate(report['sections']):
            content += f"{i+1}. {section['name']}\n\n"
            content += f"{section.get('content', 'No content available')}\n\n"
            content += "-" * 50 + "\n\n"

        return content.encode('utf-8')

    except Exception as e:
        # Ultimate fallback - just return text content
        content = f"{report['title']}\n\nGenerated on {report['generation_date']}\n\n"

        for i, section in enumerate(report['sections']):
            content += f"{i+1}. {section['name']}\n\n"
            content += f"{section.get('content', 'No content available')}\n\n"

        return content.encode('utf-8')

def create_basic_word_content_with_comments(report, include_validation_comments=True):
    """Create a Word document with validation comments using add_comment function"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create new document
        doc = Document()

        # Add title
        title = doc.add_heading(report['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add generation date
        date_para = doc.add_paragraph(f"Generated on {report['generation_date']}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page break
        doc.add_page_break()

        # Add sections with validation comments
        for i, section in enumerate(report['sections']):
            # Section heading
            section_heading = doc.add_heading(f"{i+1}. {section['name']}", level=1)

            # Section content
            content = section.get('content', 'No content available')
            validation_results = section.get('validation_results', {})

            if include_validation_comments and validation_results:
                # Add content with validation comments
                add_content_with_validation_comments(doc, content, validation_results)
            else:
                # Add content without comments
                paragraphs = content.split('\n\n')
                for paragraph in paragraphs:
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())

            # Add spacing
            doc.add_paragraph()

        # Save to BytesIO
        from io import BytesIO
        word_buffer = BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)

        return word_buffer.getvalue()

    except ImportError:
        # If python-docx is not available, create a simple text-based "Word" document
        content = f"{report['title']}\n\nGenerated on {report['generation_date']}\n\n"

        for i, section in enumerate(report['sections']):
            content += f"{i+1}. {section['name']}\n\n"
            section_content = section.get('content', 'No content available')
            content += f"{section_content}\n\n"

            # Add validation comments as text if enabled
            if include_validation_comments:
                validation_results = section.get('validation_results', {})
                if validation_results and validation_results.get('issues'):
                    content += "\n--- VALIDATION COMMENTS ---\n"
                    for issue in validation_results.get('issues', []):
                        severity = issue.get('severity', 'medium').upper()
                        issue_type = issue.get('issue_type', 'Issue').title()
                        description = issue.get('description', 'No description')
                        content += f"[{severity} - {issue_type}] {description}\n"
                    content += "--- END VALIDATION COMMENTS ---\n\n"

            content += "-" * 50 + "\n\n"

        return content.encode('utf-8')

    except Exception as e:
        # Ultimate fallback - just return text content with validation
        content = f"{report['title']}\n\nGenerated on {report['generation_date']}\n\n"

        for i, section in enumerate(report['sections']):
            content += f"{i+1}. {section['name']}\n\n"
            section_content = section.get('content', 'No content available')
            content += f"{section_content}\n\n"

        return content.encode('utf-8')

def add_content_with_validation_comments(doc, content, validation_results):
    """Add content to document with validation comments using add_comment function"""

    # Split content into paragraphs
    paragraphs = content.split('\n\n')

    # Get validation issues
    validation_issues = validation_results.get('issues', [])

    for paragraph_text in paragraphs:
        if paragraph_text.strip():
            paragraph = doc.add_paragraph()

            # Find validation issues that apply to this paragraph
            paragraph_issues = []
            for issue in validation_issues:
                text_span = issue.get('text_span', '')
                if text_span and text_span.strip() in paragraph_text:
                    paragraph_issues.append(issue)

            if paragraph_issues:
                # Add paragraph text with highlighting and comments
                add_paragraph_with_comments(paragraph, paragraph_text, paragraph_issues)
            else:
                # Add paragraph without comments
                paragraph.add_run(paragraph_text)

def add_paragraph_with_comments(paragraph, text, issues):
    """Add paragraph text with validation comments"""

    # For each issue, we'll add the text and then a comment
    current_pos = 0

    for issue in issues:
        text_span = issue.get('text_span', '').strip()

        if text_span and text_span in text:
            # Find position of the text span
            span_start = text.find(text_span, current_pos)
            if span_start != -1:
                span_end = span_start + len(text_span)

                # Add text before the span
                if span_start > current_pos:
                    paragraph.add_run(text[current_pos:span_start])

                # Add the span with highlighting
                highlighted_run = paragraph.add_run(text_span)

                # Apply highlighting based on severity
                severity = issue.get('severity', 'medium')
                if severity == 'high':
                    # Red highlight for high severity
                    highlighted_run.font.highlight_color = 2  # Red
                elif severity == 'medium':
                    # Yellow highlight for medium severity
                    highlighted_run.font.highlight_color = 7  # Yellow
                else:
                    # Green highlight for low severity
                    highlighted_run.font.highlight_color = 4  # Green

                # Use your add_comment function here
                try:
                    comment_text = create_validation_comment_text(issue)
                    add_comment(paragraph, highlighted_run, comment_text, issue.get('severity', 'medium'))
                except Exception as comment_error:
                    # Fallback: add comment as footnote text
                    comment_run = paragraph.add_run(f" [VALIDATION: {comment_text}]")
                    comment_run.font.size = Pt(8)
                    comment_run.font.italic = True

                current_pos = span_end

    # Add remaining text
    if current_pos < len(text):
        paragraph.add_run(text[current_pos:])

def create_validation_comment_text(issue):
    """Create formatted comment text from validation issue"""
    severity = issue.get('severity', 'medium').upper()
    issue_type = issue.get('issue_type', 'Issue').title()
    description = issue.get('description', 'No description available')
    suggested_fix = issue.get('suggested_fix', '')
    confidence = issue.get('confidence_score', 0.0)

    comment_text = f"[{severity} - {issue_type}]\n{description}"

    if suggested_fix:
        comment_text += f"\n\nSuggested Fix: {suggested_fix}"

    if confidence > 0:
        comment_text += f"\n\nAI Confidence: {confidence:.1%}"

    return comment_text

def add_comment(paragraph, run, comment_text, severity):
    """Add comment to Word document - placeholder for your custom function"""
    # This is where you would integrate your custom add_comment function
    # For now, this is a placeholder that adds inline comments

    # Your custom add_comment function should be called here like:
    # add_comment(paragraph, run, comment_text, severity)

    # Placeholder implementation - replace with your actual function
    try:
        # If you have python-docx-template or custom commenting, use it here
        # For now, we'll add a simple text annotation
        from docx.shared import Pt

        comment_run = paragraph.add_run(f" [üí¨ {comment_text}]")
        comment_run.font.size = Pt(8)
        comment_run.font.italic = True

        # Color code by severity
        if severity == 'high':
            comment_run.font.color.rgb = (204, 0, 0)  # Red
        elif severity == 'medium':
            comment_run.font.color.rgb = (255, 102, 0)  # Orange
        else:
            comment_run.font.color.rgb = (0, 102, 204)  # Blue

    except Exception as e:
        # Ultimate fallback
        pass

@st.fragment
def word_download_fragment(report):
    """Fragment for Word document download - prevents full page reload"""
    st.markdown("### üìÑ Word Document")

    include_validation_checked = st.checkbox("Include AI validation comments", value=True, key="word_validation_frag")

    # Initialize word preparation state
    if 'word_preparation_state' not in st.session_state:
        st.session_state.word_preparation_state = 'ready'  # ready, preparing, prepared, error
        st.session_state.word_download_data = None
        st.session_state.word_preparation_logs = []

    # Debug logging function
    def log_debug(message):
        timestamp = time.strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {message}"
        st.session_state.word_preparation_logs.append(log_entry)
        print(f"DEBUG: {log_entry}")  # Console logging

    # Show current status and controls based on state
    if st.session_state.word_preparation_state == 'ready':
        st.info("üí° Ready to prepare Word document")

        # Show debug logs if any exist
        if st.session_state.word_preparation_logs:
            with st.expander("üîç Debug Logs", expanded=False):
                for log_entry in st.session_state.word_preparation_logs[-5:]:  # Show last 5 logs
                    st.text(log_entry)

        if st.button("üìÑ Prepare Word Document", type="primary", use_container_width=True, key="prepare_word_frag"):
            log_debug("User clicked Prepare Word Document button")
            st.session_state.word_preparation_state = 'preparing'
            st.rerun()

    elif st.session_state.word_preparation_state == 'preparing':
        st.warning("üîÑ Preparing document... please wait")

        # Debug: Check if report data is available
        log_debug(f"Report data available: {report is not None}")
        if report:
            log_debug(f"Report title: {report.get('title', 'No title')}")
            log_debug(f"Report sections: {len(report.get('sections', []))}")
        else:
            log_debug("ERROR: No report data available for Word generation!")
            st.error("‚ùå No report data available for Word generation")
            st.session_state.word_preparation_state = 'error'
            st.session_state.word_error_details = {
                'advanced_error': 'No report data available',
                'basic_error': 'Report parameter is None'
            }
            st.rerun()
            return

        # Show progress indicator
        progress_bar = st.progress(0)
        status_text = st.empty()

        log_debug("Starting Word document preparation")

        try:
            status_text.text("üîÑ Step 1/4: Calling backend export service...")
            progress_bar.progress(0.1)
            log_debug("About to call create_report_and_export_word()")

            # First try advanced export through backend
            export_result = create_report_and_export_word(report, include_validation_checked)
            log_debug(f"Export result received: {export_result}")

            status_text.text("üì° Step 2/4: Processing backend response...")
            progress_bar.progress(0.3)

            if export_result and not export_result.get('error'):
                download_url = export_result.get('download_url')
                log_debug(f"Download URL: {download_url}")

                if download_url:
                    status_text.text("‚¨áÔ∏è Step 3/4: Downloading Word document...")
                    progress_bar.progress(0.6)
                    log_debug("About to call download_word_document()")

                    # Download the file content
                    word_content = download_word_document(download_url)
                    log_debug(f"Downloaded content size: {len(word_content) if word_content else 'None'}")

                    if word_content:
                        status_text.text("‚úÖ Step 4/4: Finalizing document...")
                        progress_bar.progress(0.9)

                        filename = export_result.get('filename', f"{report['title'].replace(' ', '_')}.docx")
                        st.session_state.word_download_data = {
                            'content': word_content,
                            'filename': filename,
                            'type': 'advanced'
                        }
                        st.session_state.word_preparation_state = 'prepared'
                        log_debug(f"Advanced Word document prepared successfully: {filename}")

                        status_text.text("‚úÖ Document ready!")
                        progress_bar.progress(1.0)
                    else:
                        raise Exception("Failed to download generated document")
                else:
                    raise Exception("No download URL provided")
            else:
                error_msg = export_result.get('error', 'Unknown error') if export_result else 'No response'
                raise Exception(f"Backend export failed: {error_msg}")

        except Exception as e:
            log_debug(f"Advanced export failed: {str(e)}")

            # Fallback to basic Word document
            try:
                status_text.text("üîÑ Fallback: Creating basic Word document...")
                progress_bar.progress(0.7)
                log_debug("Trying basic Word document generation")

                basic_word_content = create_basic_word_content(report)
                log_debug(f"Basic content size: {len(basic_word_content) if basic_word_content else 'None'}")

                st.session_state.word_download_data = {
                    'content': basic_word_content,
                    'filename': f"{report['title'].replace(' ', '_')}_basic.docx",
                    'type': 'basic'
                }
                st.session_state.word_preparation_state = 'prepared'
                log_debug("Basic Word document prepared successfully")

                status_text.text("‚úÖ Basic document ready!")
                progress_bar.progress(1.0)

            except Exception as basic_error:
                log_debug(f"Both advanced and basic document generation failed: {str(basic_error)}")
                st.session_state.word_preparation_state = 'error'
                st.session_state.word_error_details = {
                    'advanced_error': str(e),
                    'basic_error': str(basic_error)
                }
                status_text.text("‚ùå Document preparation failed")
                progress_bar.progress(0)

        # Auto-rerun to show the result after a brief delay
        time.sleep(2)
        st.rerun()

    elif st.session_state.word_preparation_state == 'prepared' and st.session_state.word_download_data:
        # DOWNLOAD READY STATE - stays visible after clicking download
        word_data = st.session_state.word_download_data
        doc_type = "Professional" if word_data['type'] == 'advanced' else "Basic"
        file_size = len(word_data['content'])

        st.success(f"‚úÖ {doc_type} document ready!")
        st.info(f"üìÑ {word_data['filename']} ({file_size/1024:.1f} KB)")

        # Primary download button - PERSISTENT - NO RERUN
        st.download_button(
            label=f"üìÑ Download {doc_type} Document",
            data=word_data['content'],
            file_name=word_data['filename'],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="download_word_fragment",
            help="Ready for download! Click multiple times if needed."
        )

        # Show debug logs
        if st.session_state.word_preparation_logs:
            with st.expander("üîç Debug Logs", expanded=False):
                for log_entry in st.session_state.word_preparation_logs[-10:]:  # Show last 10 logs
                    st.text(log_entry)

        # Option to prepare new document (smaller button)
        if st.button("üîÑ Prepare New", use_container_width=True, key="prepare_new_word_frag"):
            st.session_state.word_preparation_state = 'ready'
            st.session_state.word_download_data = None
            st.rerun()

    elif st.session_state.word_preparation_state == 'error':
        st.error("‚ùå Document preparation failed")

        # Show error details if available
        if hasattr(st.session_state, 'word_error_details'):
            error_details = st.session_state.word_error_details
            with st.expander("üîç Error Details", expanded=True):
                st.text("Advanced Export Error:")
                st.code(error_details.get('advanced_error', 'Unknown error'))
                st.text("Basic Export Error:")
                st.code(error_details.get('basic_error', 'Unknown error'))

        # Show debug logs
        if st.session_state.word_preparation_logs:
            with st.expander("üîç Debug Logs", expanded=True):
                for log_entry in st.session_state.word_preparation_logs:
                    st.text(log_entry)

        if st.button("üîÑ Try Again", use_container_width=True, key="retry_word_frag"):
            st.session_state.word_preparation_state = 'ready'
            st.session_state.word_download_data = None
            if hasattr(st.session_state, 'word_error_details'):
                delattr(st.session_state, 'word_error_details')
            st.rerun()

def word_download_section(report):
    """Word document download section with proper spinner and state management"""
    st.markdown("### üìÑ Word Document")

    include_validation_checked = st.checkbox("Include AI validation comments", value=True, key="word_validation_main")

    # Initialize word preparation state
    if 'word_download_data' not in st.session_state:
        st.session_state.word_download_data = None
    if 'word_preparation_logs' not in st.session_state:
        st.session_state.word_preparation_logs = []

    # Debug logging function
    def log_debug(message):
        timestamp = time.strftime("%H:%M:%S")
        log_entry = f"[{timestamp}] {message}"
        st.session_state.word_preparation_logs.append(log_entry)
        print(f"DEBUG: {log_entry}")  # Console logging

    # Check if we already have a prepared document
    if st.session_state.word_download_data:
        # DOWNLOAD READY STATE - stays visible after clicking download
        word_data = st.session_state.word_download_data
        doc_type = "Professional" if word_data['type'] == 'advanced' else "Basic"
        file_size = len(word_data['content'])

        st.success(f"‚úÖ {doc_type} document ready!")
        st.info(f"üìÑ {word_data['filename']} ({file_size/1024:.1f} KB)")

        # Primary download button - PERSISTENT - NO RERUN
        st.download_button(
            label=f"üìÑ Download {doc_type} Document",
            data=word_data['content'],
            file_name=word_data['filename'],
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="download_word_main",
            help="Ready for download! Click multiple times if needed."
        )

        # Show debug logs
        if st.session_state.word_preparation_logs:
            with st.expander("üîç Debug Logs", expanded=False):
                for log_entry in st.session_state.word_preparation_logs[-10:]:  # Show last 10 logs
                    st.text(log_entry)

        # Option to prepare new document (smaller button)
        if st.button("üîÑ Prepare New", use_container_width=True, key="prepare_new_word_main"):
            st.session_state.word_download_data = None
            st.rerun()

    else:
        # READY STATE - show preparation button
        st.info("üí° Ready to prepare Word document")

        # Show debug logs if any exist
        if st.session_state.word_preparation_logs:
            with st.expander("üîç Debug Logs", expanded=False):
                for log_entry in st.session_state.word_preparation_logs[-5:]:  # Show last 5 logs
                    st.text(log_entry)

        # The main preparation button with spinner
        if st.button("üìÑ Prepare Word Document", type="primary", use_container_width=True, key="prepare_word_main"):
            log_debug("User clicked Prepare Word Document button")
            log_debug(f"Report data available: {report is not None}")

            if not report:
                log_debug("ERROR: No report data available!")
                st.error("‚ùå No report data available for Word generation")
                return

            log_debug(f"Report title: {report.get('title', 'No title')}")
            log_debug(f"Report sections: {len(report.get('sections', []))}")

            # Use st.spinner for the entire preparation process
            with st.spinner("üîÑ Preparing Word document... Please wait"):
                log_debug("=== STARTING WORD DOCUMENT PREPARATION ===")

                try:
                    log_debug("Step 1/4: Calling backend export service...")

                    # First try advanced export through backend
                    export_result = create_report_and_export_word(report, include_validation_checked)
                    log_debug(f"Export result received: {export_result}")

                    if export_result and not export_result.get('error'):
                        download_url = export_result.get('download_url')
                        log_debug(f"Download URL: {download_url}")

                        if download_url:
                            log_debug("Step 2/4: Downloading Word document...")

                            # Download the file content
                            word_content = download_word_document(download_url)
                            log_debug(f"Downloaded content size: {len(word_content) if word_content else 'None'}")

                            if word_content:
                                log_debug("Step 3/4: Finalizing document...")

                                filename = export_result.get('filename', f"{report['title'].replace(' ', '_')}.docx")
                                st.session_state.word_download_data = {
                                    'content': word_content,
                                    'filename': filename,
                                    'type': 'advanced'
                                }
                                log_debug(f"Advanced Word document prepared successfully: {filename}")

                                # SUCCESS! Rerun to show download button
                                st.rerun()
                            else:
                                raise Exception("Failed to download generated document")
                        else:
                            raise Exception("No download URL provided")
                    else:
                        error_msg = export_result.get('error', 'Unknown error') if export_result else 'No response'
                        raise Exception(f"Backend export failed: {error_msg}")

                except Exception as e:
                    log_debug(f"Advanced export failed: {str(e)}")

                    # Fallback to basic Word document
                    try:
                        log_debug("Step 4/4: Creating basic Word document as fallback...")

                        basic_word_content = create_basic_word_content_with_comments(report, include_validation_checked)
                        log_debug(f"Basic content size: {len(basic_word_content) if basic_word_content else 'None'}")

                        st.session_state.word_download_data = {
                            'content': basic_word_content,
                            'filename': f"{report['title'].replace(' ', '_')}_basic.docx",
                            'type': 'basic'
                        }
                        log_debug("Basic Word document prepared successfully")

                        # SUCCESS! Rerun to show download button
                        st.rerun()

                    except Exception as basic_error:
                        log_debug(f"Both advanced and basic document generation failed: {str(basic_error)}")
                        st.error(f"‚ùå Document preparation failed: {str(basic_error)}")

                        # Show error details
                        with st.expander("üîç Error Details", expanded=True):
                            st.text("Advanced Export Error:")
                            st.code(str(e))
                            st.text("Basic Export Error:")
                            st.code(str(basic_error))

                        return

def prepare_word_document_immediate(report, include_validation_checked, log_debug):
    """Immediately prepare word document without state transitions"""

    log_debug("=== ENTERING prepare_word_document_immediate ===")
    log_debug(f"Report parameter: {report is not None}")
    log_debug(f"Include validation: {include_validation_checked}")

    if not report:
        log_debug("ERROR: No report data available for Word generation!")
        st.error("‚ùå No report data available for Word generation")
        st.session_state.word_preparation_state = 'error'
        st.session_state.word_error_details = {
            'advanced_error': 'No report data available',
            'basic_error': 'Report parameter is None'
        }
        return

    # Show progress indicator
    progress_placeholder = st.empty()
    status_placeholder = st.empty()

    with progress_placeholder.container():
        progress_bar = st.progress(0)
        status_text = st.empty()

        log_debug("Starting Word document preparation")

        try:
            status_text.text("üîÑ Step 1/4: Calling backend export service...")
            progress_bar.progress(0.1)
            log_debug("About to call create_report_and_export_word()")
            log_debug(f"Report title for backend call: {report.get('title', 'Unknown')}")
            log_debug(f"Report sections count for backend call: {len(report.get('sections', []))}")

            # First try advanced export through backend
            export_result = create_report_and_export_word(report, include_validation_checked)
            log_debug(f"Export result received: {export_result}")

            status_text.text("üì° Step 2/4: Processing backend response...")
            progress_bar.progress(0.3)

            if export_result and not export_result.get('error'):
                download_url = export_result.get('download_url')
                log_debug(f"Download URL: {download_url}")

                if download_url:
                    status_text.text("‚¨áÔ∏è Step 3/4: Downloading Word document...")
                    progress_bar.progress(0.6)
                    log_debug("About to call download_word_document()")

                    # Download the file content
                    word_content = download_word_document(download_url)
                    log_debug(f"Downloaded content size: {len(word_content) if word_content else 'None'}")

                    if word_content:
                        status_text.text("‚úÖ Step 4/4: Finalizing document...")
                        progress_bar.progress(0.9)

                        filename = export_result.get('filename', f"{report['title'].replace(' ', '_')}.docx")
                        st.session_state.word_download_data = {
                            'content': word_content,
                            'filename': filename,
                            'type': 'advanced'
                        }
                        st.session_state.word_preparation_state = 'prepared'
                        log_debug(f"Advanced Word document prepared successfully: {filename}")

                        status_text.text("‚úÖ Document ready!")
                        progress_bar.progress(1.0)
                    else:
                        raise Exception("Failed to download generated document")
                else:
                    raise Exception("No download URL provided")
            else:
                error_msg = export_result.get('error', 'Unknown error') if export_result else 'No response'
                raise Exception(f"Backend export failed: {error_msg}")

        except Exception as e:
            log_debug(f"Advanced export failed: {str(e)}")

            # Fallback to basic Word document
            try:
                status_text.text("üîÑ Fallback: Creating basic Word document...")
                progress_bar.progress(0.7)
                log_debug("Trying basic Word document generation")

                basic_word_content = create_basic_word_content_with_comments(report, include_validation_checked)
                log_debug(f"Basic content size: {len(basic_word_content) if basic_word_content else 'None'}")

                st.session_state.word_download_data = {
                    'content': basic_word_content,
                    'filename': f"{report['title'].replace(' ', '_')}_basic.docx",
                    'type': 'basic'
                }
                st.session_state.word_preparation_state = 'prepared'
                log_debug("Basic Word document prepared successfully")

                status_text.text("‚úÖ Basic document ready!")
                progress_bar.progress(1.0)

            except Exception as basic_error:
                log_debug(f"Both advanced and basic document generation failed: {str(basic_error)}")
                st.session_state.word_preparation_state = 'error'
                st.session_state.word_error_details = {
                    'advanced_error': str(e),
                    'basic_error': str(basic_error)
                }
                status_text.text("‚ùå Document preparation failed")
                progress_bar.progress(0)

        # Clear progress after completion
        time.sleep(2)
        progress_placeholder.empty()
        status_placeholder.empty()

@st.fragment
def text_download_fragment(report):
    """Fragment for Text document download - prevents full page reload"""
    st.markdown("### üìù Text Format")

    report_text = f"# {report['title']}\n\n"
    report_text += f"Generated on {report['generation_date']}\n\n"

    for i, section in enumerate(report['sections']):
        report_text += f"## {i+1}. {section['name']}\n\n"
        report_text += f"{section['content']}\n\n"

    text_size = len(report_text.encode('utf-8'))
    st.info(f"üìÑ Text file ({text_size/1024:.1f} KB)")

    st.download_button(
        label="üìù Download as Text",
        data=report_text,
        file_name=f"{report['title'].replace(' ', '_')}.txt",
        mime="text/plain",
        use_container_width=True,
        help="Plain text format - works everywhere"
    )

@st.fragment
def json_download_fragment(report):
    """Fragment for JSON document download - prevents full page reload"""
    st.markdown("### üìã JSON Format")

    report_json = json.dumps(report, indent=2)
    json_size = len(report_json.encode('utf-8'))
    st.info(f"üìÑ JSON file ({json_size/1024:.1f} KB)")

    st.download_button(
        label="üìã Download as JSON",
        data=report_json,
        file_name=f"{report['title'].replace(' ', '_')}.json",
        mime="application/json",
        use_container_width=True,
        help="Structured data format - includes validation details"
    )

def main():
    st.title("üìä Credit Review Document System")
    
    # Check backend connectivity
    if not check_backend_health():
        st.error("üî¥ Backend server is not running. Please start the backend first.")
        st.code("cd backend && uvicorn app.main:app --reload")
        return
    
    st.success("üü¢ Backend server connected")
    
    # Simple page navigation
    page = st.sidebar.radio("Navigation", ["üìÑ Document Upload", "üìù Generate Report"])
    
    if page == "üìÑ Document Upload":
        show_upload_page()
    else:
        show_report_generation_page()

def show_upload_page():
    """Simplified document upload page"""
    st.header("üìÑ Document Upload")
    st.write("Upload your financial documents, contracts, and reports for analysis.")
    
    # File upload section
    uploaded_files = st.file_uploader(
        "Select documents to upload",
        accept_multiple_files=True,
        type=['pdf', 'docx', 'xlsx', 'jpg', 'jpeg', 'png'],
        help="Supported formats: PDF, Word, Excel, Images"
    )
    
    if uploaded_files:
        st.write(f"**{len(uploaded_files)} file(s) selected:**")
        
        # Show file details
        for file in uploaded_files:
            st.write(f"‚Ä¢ {file.name} ({file.size / 1024:.1f} KB)")
        
        if st.button("üöÄ Upload & Process All Documents", type="primary", use_container_width=True):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for i, uploaded_file in enumerate(uploaded_files):
                status_text.text(f"Processing {uploaded_file.name}...")
                
                result = upload_document(uploaded_file)
                
                if result:
                    st.success(f"‚úÖ {uploaded_file.name} processed successfully!")
                else:
                    st.error(f"‚ùå Failed to process {uploaded_file.name}")
                
                progress_bar.progress((i + 1) / len(uploaded_files))
            
            status_text.text("Processing complete!")
            st.rerun()  # Refresh to show new documents

    # Separator
    st.markdown("---")

    # Ultimate Fallback Processing Section
    st.header("üß† AI-Powered Document Processing (Sonnet Fallback)")
    st.write("When traditional document processing fails or you're not satisfied with results, use our AI-powered fallback with Claude Sonnet.")

    with st.expander("üöÄ Ultimate AI Fallback Processing", expanded=False):
        st.markdown("""
        **üåü Excellent for:**
        - üìÑ **Multi-page PDF files** (all pages auto-converted to images for AI processing)
        - üì∑ Scanned document images (photos of pages)
        - üñºÔ∏è Screenshots of documents or reports
        - üìã Images with text, tables, or charts
        - üìÑ Plain text files needing AI enhancement

        **üìã PDF Processing:**
        - Processes up to 20 pages automatically
        - Each page converted to high-quality image
        - Page markers added for reference
        - Maintains document structure across pages

        **‚ö†Ô∏è Limited support:**
        - Word documents (save as PDF first)
        """)

        col1, col2 = st.columns([3, 1])

        with col1:
            # File upload for Sonnet processing
            sonnet_uploaded_file = st.file_uploader(
                "Select document for AI processing",
                type=['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'txt', 'pdf'],
                help="üåü EXCELLENT: Images (JPG, PNG) & PDFs | Good: Text files",
                key="sonnet_uploader"
            )

            # Document name
            if sonnet_uploaded_file:
                sonnet_document_name = st.text_input(
                    "Document Name (Optional)",
                    value=sonnet_uploaded_file.name.rsplit('.', 1)[0],
                    help="Custom name for the document",
                    key="sonnet_doc_name"
                )

        with col2:
            # Processing options
            st.markdown("**‚öôÔ∏è Processing Options**")

            processing_mode = st.selectbox(
                "Processing Mode",
                options=["comprehensive", "financial", "legal"],
                help="""
                ‚Ä¢ Comprehensive: Extract all content
                ‚Ä¢ Financial: Focus on financial data
                ‚Ä¢ Legal: Preserve legal language
                """,
                key="sonnet_mode"
            )

            focus_areas = st.text_input(
                "Focus Areas (Optional)",
                placeholder="e.g., revenue, risks, compliance",
                help="Comma-separated areas to focus on",
                key="sonnet_focus"
            )

        # File size and format validation
        if sonnet_uploaded_file:
            file_size_mb = sonnet_uploaded_file.size / (1024 * 1024)

            if file_size_mb > 10:
                st.error(f"‚ùå File too large: {file_size_mb:.1f}MB (max: 10MB)")
            else:
                st.info(f"üìÑ **{sonnet_uploaded_file.name}** ({file_size_mb:.1f}MB) - Ready for AI processing")

                # Processing options info
                st.markdown("**üîç AI Processing Details:**")
                col_info1, col_info2, col_info3 = st.columns(3)

                with col_info1:
                    st.write("**Engine:** Claude Sonnet")
                    st.write("**Method:** Multimodal AI")

                with col_info2:
                    st.write("**Quality:** High precision")
                    st.write("**Speed:** ~30-60 sec/page")

                with col_info3:
                    st.write("**Formats:** All supported")
                    st.write("**Images:** OCR + Understanding")

                # Process button
                if st.button("üß† Process with AI (Sonnet)", type="primary", use_container_width=True, key="sonnet_process"):
                    if file_size_mb <= 10:
                        with st.spinner("ü§ñ Claude Sonnet is analyzing your document... Multi-page PDFs may take several minutes..."):

                            # Show processing steps
                            progress_container = st.container()
                            with progress_container:
                                step_progress = st.empty()
                                step_progress.info("üîç **Step 1/4:** Uploading document to Sonnet...")
                                time.sleep(2)

                                step_progress.info("üìñ **Step 2/4:** AI is reading and understanding content (processing each page)...")
                                time.sleep(5)

                                step_progress.info("‚úÇÔ∏è **Step 3/4:** Extracting and structuring text...")
                                time.sleep(2)

                                step_progress.info("üß© **Step 4/4:** Creating RAG chunks...")

                                # Process with Sonnet
                                result = upload_document_with_sonnet_fallback(
                                    file=sonnet_uploaded_file,
                                    document_name=sonnet_document_name,
                                    processing_mode=processing_mode,
                                    focus_areas=focus_areas if focus_areas else None
                                )

                                step_progress.empty()

                            if result:
                                st.success(f"üéâ **AI Processing Successful!**")

                                # Show detailed results
                                col_result1, col_result2, col_result3 = st.columns(3)

                                with col_result1:
                                    st.metric("Chunks Created", result.get('chunks_created', 0))

                                with col_result2:
                                    text_length = result.get('extracted_text_length', 0)
                                    st.metric("Text Extracted", f"{text_length:,} chars")

                                with col_result3:
                                    st.metric("Processing Method", "Sonnet AI")

                                # Success message with details
                                st.markdown(f"""
                                **üìä Processing Results:**
                                - **Document ID:** `{result.get('document_id', 'Unknown')}`
                                - **Chunks Generated:** {result.get('chunks_created', 0)}
                                - **Text Length:** {result.get('extracted_text_length', 0):,} characters
                                - **Processing Mode:** {processing_mode.title()}
                                - **Status:** Ready for RAG queries

                                ‚úÖ Your document has been successfully processed with Claude Sonnet and is now available in the RAG system!
                                """)

                                time.sleep(2)
                                st.rerun()  # Refresh to show new document

                            else:
                                st.error("‚ùå **AI Processing Failed**")
                                st.write("The AI processing encountered an error. Please try:")
                                st.write("‚Ä¢ Checking if the document is readable")
                                st.write("‚Ä¢ Reducing file size if possible")
                                st.write("‚Ä¢ Using a different format (PDF ‚Üí image)")
                                st.write("‚Ä¢ Contacting support if issues persist")
                    else:
                        st.error("‚ùå Please select a file under 10MB")

        # Help section
        st.markdown("---")
        st.markdown("**üí° Tips for Best Results:**")
        col_tip1, col_tip2 = st.columns(2)

        with col_tip1:
            st.write("**For Financial Documents:**")
            st.write("‚Ä¢ Use 'financial' processing mode")
            st.write("‚Ä¢ Focus areas: revenue, profit, cash flow")
            st.write("‚Ä¢ PDF format preferred")

        with col_tip2:
            st.write("**For Document Images:**")
            st.write("‚Ä¢ Take clear, high-resolution photos")
            st.write("‚Ä¢ Ensure good lighting and contrast")
            st.write("‚Ä¢ Use 'comprehensive' mode for mixed content")

    st.markdown("---")

    # Public Company Document Fetching Section
    st.header("üè¢ Public Company Document Fetching")
    st.write("Automatically fetch SEC EDGAR filings and international regulatory documents for public companies.")

    with st.expander("üìä Fetch Public Company Documents", expanded=False):
        col1, col2 = st.columns([2, 1])

        with col1:
            # Ticker symbol input
            ticker_symbol = st.text_input(
                "üî§ Ticker Symbol / Security Identifier",
                placeholder="e.g., AAPL, TSLA, MSFT",
                help="Enter the company's ticker symbol or security identifier"
            )

            # Exchange selection
            exchange = st.selectbox(
                "üåç Exchange",
                options=["US", "UK", "EU", "Canada"],
                help="Select the main exchange where the company is traded"
            )

        with col2:
            # Year and quarter selection
            current_year = time.strftime("%Y")
            year = st.selectbox(
                "üìÖ Year",
                options=[str(y) for y in range(int(current_year), int(current_year) - 5, -1)],
                help="Select the reporting year"
            )

            if exchange == "US":
                quarter = st.selectbox(
                    "üìä Quarter (Optional)",
                    options=["", "Q1", "Q2", "Q3", "Q4"],
                    help="Leave empty for annual reports, select quarter for 10-Q filings"
                )
            else:
                quarter = st.selectbox(
                    "üìä Period",
                    options=["Annual", "Half-Year", "Q1", "Q2", "Q3", "Q4"],
                    help="Select the reporting period"
                )

        # Filing type selection
        if exchange == "US":
            st.markdown("üìã **SEC Filing Types to Fetch:**")
            col_10k, col_10q, col_8k, col_proxy = st.columns(4)

            with col_10k:
                fetch_10k = st.checkbox("10-K (Annual)", value=True, help="Annual report")
            with col_10q:
                fetch_10q = st.checkbox("10-Q (Quarterly)", value=True, help="Quarterly report")
            with col_8k:
                fetch_8k = st.checkbox("8-K (Current)", help="Current report")
            with col_proxy:
                fetch_proxy = st.checkbox("DEF 14A (Proxy)", help="Proxy statement")
        else:
            st.markdown("üìã **Document Types to Fetch:**")
            col_annual, col_interim, col_governance = st.columns(3)

            with col_annual:
                fetch_annual = st.checkbox("Annual Report", value=True)
            with col_interim:
                fetch_interim = st.checkbox("Interim Report", value=True)
            with col_governance:
                fetch_governance = st.checkbox("Governance Docs", help="Corporate governance documents")

        # Additional options
        st.markdown("‚öôÔ∏è **Additional Options:**")
        col_opt1, col_opt2 = st.columns(2)

        with col_opt1:
            include_exhibits = st.checkbox("Include Exhibits", help="Fetch document exhibits when available")
        with col_opt2:
            auto_process = st.checkbox("Auto-process for RAG", value=True, help="Automatically process documents for RAG system")

        # Fetch button
        if st.button("üöÄ Fetch Public Company Documents", type="primary", use_container_width=True):
            # Validation
            if not ticker_symbol:
                st.error("‚ùå Please enter a ticker symbol")
            elif not ticker_symbol.replace("-", "").replace(".", "").isalnum():
                st.error("‚ùå Invalid ticker symbol format")
            else:
                # Prepare filing types based on exchange
                filing_types = []
                if exchange == "US":
                    if fetch_10k:
                        filing_types.append("10-K")
                    if fetch_10q:
                        filing_types.append("10-Q")
                    if fetch_8k:
                        filing_types.append("8-K")
                    if fetch_proxy:
                        filing_types.append("DEF 14A")
                else:
                    if locals().get('fetch_annual', False):
                        filing_types.append("Annual Report")
                    if locals().get('fetch_interim', False):
                        filing_types.append("Interim Report")
                    if locals().get('fetch_governance', False):
                        filing_types.append("Governance")

                if not filing_types:
                    st.error("‚ùå Please select at least one document type to fetch")
                else:
                    # Show progress
                    with st.spinner(f"üîç Fetching documents for {ticker_symbol.upper()} ({exchange})..."):
                        result = fetch_public_company_documents(
                            ticker_symbol=ticker_symbol.upper(),
                            exchange=exchange,
                            quarter=quarter if quarter else None,
                            year=int(year),
                            filing_types=filing_types
                        )

                    if result:
                        if result.get("success"):
                            documents_found = result.get("documents", [])
                            st.success(f"‚úÖ Successfully fetched {len(documents_found)} document(s) for {ticker_symbol.upper()}")

                            # Show summary of fetched documents
                            if documents_found:
                                st.markdown("üìÑ **Fetched Documents:**")
                                for doc in documents_found:
                                    st.write(f"‚Ä¢ **{doc.get('filing_type', 'Unknown')}** - {doc.get('title', 'Untitled')} ({doc.get('date', 'Unknown date')})")

                                if auto_process:
                                    st.info("üîÑ Documents are being automatically processed for RAG system...")
                                    time.sleep(2)  # Give time for processing
                                    st.rerun()  # Refresh to show new documents
                            else:
                                st.warning("‚ö†Ô∏è No documents found for the specified criteria")
                        else:
                            error_msg = result.get("error", "Unknown error occurred")
                            st.error(f"‚ùå Failed to fetch documents: {error_msg}")
                    else:
                        st.error("‚ùå Failed to connect to document fetching service")

    st.markdown("---")

    # Document library
    st.subheader("üìö Document Library")
    documents = get_documents()
    
    if documents:
        # Create a nice display of documents with delete buttons
        st.subheader("üìö Document Library")
        
        for i, doc in enumerate(documents):
            doc_id = doc.get('id', doc.get('document_id'))
            doc_name = doc.get('name', 'Unknown')
            chunk_count = doc.get('chunk_count', 0)
            upload_date = doc.get('upload_date', 'Unknown')[:16] if doc.get('upload_date') else 'Unknown'
            file_size = doc.get('size', 0)
            
            # Create expandable container for each document
            with st.expander(f"üìÑ {doc_name} ({chunk_count} chunks)", expanded=False):
                col1, col2, col3 = st.columns([3, 1, 1])
                
                with col1:
                    st.write(f"**ID:** {doc_id}")
                    st.write(f"**Chunks:** {chunk_count}")
                    st.write(f"**Upload Date:** {upload_date}")
                    st.write(f"**Size:** {file_size / 1024:.1f} KB" if file_size > 0 else "**Size:** Unknown")
                    
                    # Show special indicator for 0-chunk documents
                    if chunk_count == 0:
                        st.warning("‚ö†Ô∏è This document has 0 chunks - it may have failed to process or contains only images")
                
                with col2:
                    if st.button(f"üîç View Details", key=f"view_{doc_id}", use_container_width=True):
                        # Store the document ID to show details
                        st.session_state[f"show_details_{doc_id}"] = True
                        st.rerun()

                # Show document details including chunk samples
                if st.session_state.get(f"show_details_{doc_id}", False):
                    with st.container():
                        st.markdown(f"### üìÑ Details for {doc_name}")

                        # Basic info
                        col_info1, col_info2 = st.columns(2)
                        with col_info1:
                            st.write(f"**Document ID:** `{doc_id}`")
                            st.write(f"**Total Chunks:** {chunk_count}")
                        with col_info2:
                            st.write(f"**Upload Date:** {upload_date}")
                            st.write(f"**File Size:** {file_size / 1024:.1f} KB" if file_size > 0 else "**File Size:** Unknown")

                        # Show chunk samples
                        if chunk_count > 0:
                            st.markdown("#### üìù Random Chunk Samples")
                            st.write("Here are 2 randomly selected chunks to help you understand the document content:")

                            with st.spinner("Loading chunk samples..."):
                                chunks_data = get_document_chunks(doc_id, limit=2)

                            if chunks_data and chunks_data.get("chunks"):
                                for i, chunk in enumerate(chunks_data["chunks"]):
                                    chunk_type = chunk.get("type", "unknown")
                                    chunk_content = chunk.get("content", "No content")
                                    chunk_metadata = chunk.get("metadata", {})

                                    # Display chunk info in a bordered container instead of expander
                                    st.markdown(f"**üìÑ Sample {i+1}: {chunk_type.title()} Chunk**")

                                    # Create a bordered container using markdown
                                    with st.container():
                                        # Show metadata if available
                                        if chunk_metadata:
                                            metadata_col, content_col = st.columns([1, 2])
                                            with metadata_col:
                                                st.markdown("**üìã Metadata:**")
                                                for key, value in chunk_metadata.items():
                                                    if key not in ['embedding']:  # Skip embeddings
                                                        st.write(f"**{key.title()}:** {value}")

                                            with content_col:
                                                st.markdown("**üìù Content:**")
                                                # Truncate content if too long
                                                display_content = chunk_content
                                                if len(display_content) > 500:
                                                    display_content = display_content[:500] + "..."
                                                st.text_area("", value=display_content, height=150, disabled=True, key=f"chunk_content_{doc_id}_{i}")
                                        else:
                                            st.markdown("**üìù Content:**")
                                            display_content = chunk_content
                                            if len(display_content) > 500:
                                                display_content = display_content[:500] + "..."
                                            st.text_area("", value=display_content, height=150, disabled=True, key=f"chunk_content_{doc_id}_{i}")

                                    # Add visual separator
                                    if i < len(chunks_data["chunks"]) - 1:
                                        st.markdown("---")
                            else:
                                st.warning("‚ùå Unable to load chunk samples. This might indicate processing issues.")
                        else:
                            st.warning("‚ö†Ô∏è This document has no chunks available for preview.")

                        # Close button
                        if st.button("‚ùå Close Details", key=f"close_details_{doc_id}"):
                            st.session_state[f"show_details_{doc_id}"] = False
                            st.rerun()

                        st.markdown("---")
                
                with col3:
                    if st.button(f"üóëÔ∏è Delete", key=f"delete_{doc_id}", type="secondary", use_container_width=True):
                        # Confirmation dialog using session state
                        st.session_state[f"confirm_delete_{doc_id}"] = True
                
                # Handle delete confirmation
                if st.session_state.get(f"confirm_delete_{doc_id}", False):
                    st.error(f"‚ö†Ô∏è Are you sure you want to delete '{doc_name}'?")
                    col_confirm, col_cancel = st.columns(2)
                    
                    with col_confirm:
                        if st.button("‚úÖ Yes, Delete", key=f"confirm_yes_{doc_id}", type="primary"):
                            with st.spinner(f"Deleting {doc_name}..."):
                                result = delete_document(doc_id)
                                if result:
                                    st.success(f"‚úÖ {doc_name} deleted successfully!")
                                    st.session_state[f"confirm_delete_{doc_id}"] = False
                                    time.sleep(1)
                                    st.rerun()
                                else:
                                    st.error(f"‚ùå Failed to delete {doc_name}")
                    
                    with col_cancel:
                        if st.button("‚ùå Cancel", key=f"confirm_no_{doc_id}"):
                            st.session_state[f"confirm_delete_{doc_id}"] = False
                            st.rerun()
        
        # Also show as a summary table
        st.subheader("üìä Document Summary")
        doc_df = pd.DataFrame([{
            "Document Name": doc.get("name", "Unknown"),
            "Chunks": doc.get("chunk_count", 0),
            "Status": "‚ö†Ô∏è Needs Review" if doc.get("chunk_count", 0) == 0 else "‚úÖ Processed",
            "Size (KB)": f"{doc.get('size', 0) / 1024:.1f}" if doc.get('size', 0) > 0 else "Unknown",
            "Upload Date": doc.get("upload_date", "Unknown")[:16] if doc.get("upload_date") else "Unknown"
        } for doc in documents])
        
        st.dataframe(doc_df, use_container_width=True)
        
        # Quick stats
        total_docs = len(documents)
        zero_chunk_docs = len([d for d in documents if d.get('chunk_count', 0) == 0])
        total_chunks = sum(d.get('chunk_count', 0) for d in documents)
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Documents", total_docs)
        with col2:
            st.metric("Total Chunks", total_chunks)
        with col3:
            st.metric("Failed Processing", zero_chunk_docs)
            if zero_chunk_docs > 0:
                st.caption("‚ö†Ô∏è Documents with 0 chunks may need reprocessing")
    else:
        st.info("üì≠ No documents uploaded yet. Use the uploader above to get started.")

def show_report_generation_page():
    """Unified report generation page"""
    st.header("üìù Generate Report")
    st.write("Select documents and create report segments with AI assistance.")
    
    # Get available documents
    documents = get_documents()
    
    if not documents:
        st.warning("‚ö†Ô∏è No documents available. Please upload documents first.")
        if st.button("Go to Document Upload"):
            st.session_state.page = "üìÑ Document Upload"
            st.rerun()
        return
    
    # Document selection
    st.subheader("üìã 1. Select Documents")
    
    # Show available documents with checkboxes
    selected_docs = []
    
    with st.expander("üìö Available Documents", expanded=True):
        select_all = st.checkbox("Select All Documents")
        
        for i, doc in enumerate(documents):
            is_selected = select_all or st.checkbox(
                f"{doc.get('name', 'Unknown')} ({doc.get('chunk_count', 0)} chunks)",
                key=f"doc_{i}"
            )
            if is_selected:
                selected_docs.append(doc)
    
    if not selected_docs:
        st.info("üëÜ Please select at least one document above.")
        return
    
    st.success(f"‚úÖ {len(selected_docs)} document(s) selected")
    
    # Report segments
    st.subheader("üìù 2. Report Segments")
    
    # Initialize segments in session state
    if "segments" not in st.session_state:
        st.session_state.segments = [
            {"name": "Executive Summary", "prompt": "Provide a high-level executive summary of the financial position and key findings."},
            {"name": "Financial Analysis", "prompt": "Analyze revenue trends, profitability, cash flow, and key financial ratios."},
            {"name": "Risk Assessment", "prompt": "Identify and assess key business risks, financial risks, and credit risks."}
        ]
    
    # Add/Edit segments
    with st.expander("‚ûï Manage Report Sections", expanded=False):
        # Add new segment
        col1, col2 = st.columns([3, 1])
        with col1:
            new_name = st.text_input("Section Name", placeholder="e.g., Market Analysis")
            new_prompt = st.text_area("Section Prompt", placeholder="Describe what this section should contain...", height=100)
        with col2:
            st.write("")  # Spacing
            if st.button("Add Section"):
                if new_name and new_prompt:
                    st.session_state.segments.append({"name": new_name, "prompt": new_prompt})
                    st.success("Section added!")
                    st.rerun()
        
        # Edit existing segments
        for i, segment in enumerate(st.session_state.segments):
            with st.container():
                st.markdown(f"**Section {i+1}:**")
                col1, col2, col3 = st.columns([2, 3, 1])
                with col1:
                    segment["name"] = st.text_input(f"Name", value=segment["name"], key=f"name_{i}")
                with col2:
                    segment["prompt"] = st.text_area(f"Prompt", value=segment["prompt"], key=f"prompt_{i}", height=80)
                with col3:
                    st.write("")  # Spacing
                    if st.button("Remove", key=f"remove_{i}"):
                        st.session_state.segments.pop(i)
                        st.rerun()
                st.markdown("---")
    
    # Show current segments
    st.write("**Current Report Sections:**")
    for i, segment in enumerate(st.session_state.segments):
        with st.expander(f"{i+1}. {segment['name']}", expanded=False):
            st.write(f"**Prompt:** {segment['prompt']}")
    
    # Generate report
    st.subheader("üöÄ 3. Generate Report")
    
    col1, col2 = st.columns([2, 1])
    with col1:
        report_title = st.text_input("Report Title", value="Credit Review Report", placeholder="Enter report title...")
    with col2:
        st.write("")  # Spacing
        generate_button = st.button("üéØ Generate Full Report", type="primary", use_container_width=True)
    
    if generate_button:
        if not report_title:
            st.error("Please enter a report title.")
            return
        
        if not st.session_state.segments:
            st.error("Please add at least one report section.")
            return
        
        # Generate report
        with st.spinner("ü§ñ Generating report sections..."):
            generated_report = generate_complete_report(
                title=report_title,
                segments=st.session_state.segments,
                selected_documents=selected_docs
            )
        
        if generated_report:
            show_generated_report(generated_report)

def generate_complete_report(title, segments, selected_documents):
    """Generate complete report with all segments"""
    
    report_sections = []
    total_sections = len(segments)
    
    # Extract document IDs from selected documents
    selected_document_ids = [doc.get('id', doc.get('document_id')) for doc in selected_documents if doc.get('id') or doc.get('document_id')]
    
    # Progress tracking
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, segment in enumerate(segments):
        status_text.text(f"Generating: {segment['name']}...")
        
        # Generate content for this segment
        try:
            # Create segment data
            segment_data = {
                "name": segment['name'],
                "prompt": segment['prompt'],
                "required_document_types": [],
                "generation_settings": {}
            }
            
            result = generate_segment_content(segment_data, selected_document_ids)
            
            if result and not result.get('error'):
                content = result.get('generated_content', 'No content generated')
                validation_results = result.get('validation_results', {})
                generation_metadata = result.get('generation_metadata', {})
                
                # Store validation results both in metadata and as a top-level field
                report_sections.append({
                    "name": segment['name'],
                    "content": content,
                    "validation_results": validation_results,
                    "metadata": {
                        **generation_metadata,
                        "validation_results": validation_results
                    }
                })
                st.success(f"‚úÖ {segment['name']} completed")
            else:
                error_msg = result.get('error', 'Unknown error') if result else 'No response from server'
                st.error(f"‚ùå Failed to generate {segment['name']}: {error_msg}")
                report_sections.append({
                    "name": segment['name'],
                    "content": f"*Content generation failed: {error_msg}*",
                    "metadata": {}
                })
        
        except Exception as e:
            st.error(f"‚ùå Error generating {segment['name']}: {str(e)}")
            report_sections.append({
                "name": segment['name'],
                "content": f"*Error: {str(e)}*",
                "metadata": {}
            })
        
        progress_bar.progress((i + 1) / total_sections)
    
    status_text.text("Report generation complete!")
    
    return {
        "title": title,
        "sections": report_sections,
        "documents_used": selected_documents,
        "generation_date": time.strftime("%Y-%m-%d %H:%M:%S")
    }

def get_paragraph_validation(validation_data, paragraph_text):
    """Get validation status for a specific paragraph"""
    # Backend returns 'issues', not 'validation_issues'
    if not validation_data or not validation_data.get('issues'):
        return {
            'status': 'passed',
            'issues': [],
            'conclusion': 'No issues detected'
        }
    
    validation_issues = validation_data.get('issues', [])
    paragraph_issues = []
    
    # Find issues that relate to this paragraph
    for issue in validation_issues:
        text_span = issue.get('text_span', '')
        if text_span and text_span.strip() in paragraph_text:
            paragraph_issues.append(issue)
    
    # Determine overall status
    if not paragraph_issues:
        status = 'passed'
        conclusion = 'All validations passed'
    else:
        high_severity_issues = [i for i in paragraph_issues if i.get('severity') == 'high']
        medium_severity_issues = [i for i in paragraph_issues if i.get('severity') == 'medium']
        
        if high_severity_issues:
            status = 'not_passed'
            conclusion = f"{len(high_severity_issues)} high-severity issues found"
        elif medium_severity_issues:
            status = 'partially_passed'
            conclusion = f"{len(medium_severity_issues)} medium-severity issues found"
        else:
            status = 'partially_passed'
            conclusion = f"{len(paragraph_issues)} minor issues found"
    
    return {
        'status': status,
        'issues': paragraph_issues,
        'conclusion': conclusion
    }

def display_paragraph_validation_status(validation_status):
    """Display paragraph validation status with detailed explanations"""
    status = validation_status.get('status', 'unknown')
    conclusion = validation_status.get('conclusion', 'Unknown status')
    issues = validation_status.get('issues', [])
    
    # Enhanced status display with detailed explanations
    if status == 'passed':
        st.success(f"‚úÖ **Validation Passed**: {conclusion}")
        # Show success details in a clean container instead of expander
        st.markdown("**Validation Checks Passed:**")
        st.write("‚úì **Accuracy Check**: Content appears factually correct based on source documents")
        st.write("‚úì **Completeness Check**: All required information elements are present")
        st.write("‚úì **Consistency Check**: No contradictory statements detected")
        st.write("‚úì **Compliance Check**: Meets regulatory and style requirements")
            
    elif status == 'partially_passed':
        st.warning(f"‚ö†Ô∏è **Validation Partially Passed**: {conclusion}")
        st.write("**‚ö†Ô∏è This paragraph has some validation concerns but is generally acceptable.**")
        st.write("Review the issues below and consider revisions to improve quality.")
            
    elif status == 'not_passed':
        st.error(f"‚ùå **Validation Failed**: {conclusion}")
        st.write("**üö® This paragraph has significant validation failures.**")
        st.write("**Immediate action required** - review and revise before using in final report.")
            
    else:
        st.info(f"‚ÑπÔ∏è **Validation Status**: {conclusion}")
    
    # Show detailed issues for this paragraph with enhanced explanations
    if issues:
        st.markdown("**üìã Validation Issues Found:**")
        for i, issue in enumerate(issues, 1):
            severity = issue.get('severity', 'medium')
            issue_type = issue.get('issue_type', 'Issue')
            description = issue.get('description', 'No description available')
            suggested_fix = issue.get('suggested_fix', '')
            confidence = issue.get('confidence_score', 0.0)
            text_span = issue.get('text_span', '')
            
            severity_icon = {
                'high': 'üî¥',
                'medium': 'üü°',
                'low': 'üü¢',
                'info': 'üîµ'
            }.get(severity, '‚ö™')
            
            severity_label = {
                'high': '**HIGH PRIORITY**',
                'medium': '**MEDIUM PRIORITY**',
                'low': '**LOW PRIORITY**',
                'info': '**INFORMATIONAL**'
            }.get(severity, '**UNKNOWN**')
            
            # Create a bordered container for each issue instead of expander
            st.markdown(f"**{severity_icon} Issue #{i}: {issue_type.title()} - {severity_label}**")
            
            # Issue details in a clean format
            col1, col2 = st.columns([1, 1])
            with col1:
                st.write(f"**Type:** {issue_type.title()}")
                st.write(f"**Severity:** {severity_label}")
            with col2:
                st.write(f"**Confidence:** {confidence:.1%}")
                if text_span:
                    st.write(f"**Affected Text:** _{text_span[:50]}..._" if len(text_span) > 50 else f"**Affected Text:** _{text_span}_")
            
            # Issue description
            st.write(f"**Description:** {description}")
            
            # Detailed explanation based on issue type
            if issue_type == 'accuracy':
                st.write("üìç **Root Cause:** Information may not be supported by source documents")
                st.write("‚ö†Ô∏è **Impact:** Could mislead readers or damage report credibility")
                
            elif issue_type == 'completeness':
                st.write("üìç **Root Cause:** Required elements may not be fully addressed")
                st.write("‚ö†Ô∏è **Impact:** May leave readers with incomplete understanding")
                
            elif issue_type == 'consistency':
                st.write("üìç **Root Cause:** Contradictory statements or inconsistent data")
                st.write("‚ö†Ô∏è **Impact:** Creates confusion and reduces report reliability")
                
            elif issue_type == 'compliance':
                st.write("üìç **Root Cause:** Missing disclosures or inappropriate language")
                st.write("‚ö†Ô∏è **Impact:** Could result in regulatory issues")
                
            elif issue_type == 'factual_error':
                st.write("üìç **Root Cause:** Misinterpretation of source documents")
                st.write("‚ö†Ô∏è **Impact:** Direct misinformation affecting decision-making")
                
            elif issue_type == 'source_mismatch':
                st.write("üìç **Root Cause:** Claims not aligned with source documents")
                st.write("‚ö†Ô∏è **Impact:** Undermines evidence-based analysis")
            
            # Suggested fix with actionable steps
            if suggested_fix:
                st.write(f"üí° **Recommended Action:** {suggested_fix}")
                
                # Additional actionable steps based on severity
                if severity == 'high':
                    st.write("üö® **Immediate Steps:** STOP ‚Üí VERIFY ‚Üí REVISE ‚Üí VALIDATE")
                elif severity == 'medium':
                    st.write("‚ö†Ô∏è **Recommended:** REVIEW ‚Üí CONSIDER ‚Üí IMPROVE")
                else:
                    st.write("‚ÑπÔ∏è **Optional:** NOTE ‚Üí CONSIDER minor refinements")
            
            st.markdown("---")

def show_generated_report(report):
    """Display the generated report with validation analysis"""
    st.subheader("üìã Dashboard - Report & Validation Analysis")
    
    # Store report in session state to persist after downloads
    st.session_state.current_report = report
    
    # Add custom CSS for better text visibility
    st.markdown("""
    <style>
    .stTextArea textarea {
        background-color: #F0F2F6 !important;
        color: #262730 !important;
        border: 1px solid #CCCCCC !important;
        font-family: 'Source Code Pro', monospace !important;
    }
    .stTextArea label {
        color: #262730 !important;
        font-weight: bold !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Report header
    st.markdown(f"# {report['title']}")
    st.caption(f"Generated on {report['generation_date']}")
    
    # Documents used
    with st.expander("üìö Documents Used", expanded=False):
        for doc in report['documents_used']:
            st.write(f"‚Ä¢ {doc.get('name', 'Unknown')} ({doc.get('chunk_count', 0)} chunks)")
    
    # Single-column layout with side-by-side paragraph containers
    for i, section in enumerate(report['sections']):
        st.markdown(f"## {i+1}. {section['name']}")
        
        # Get validation data for this section (if available)
        validation_data = section.get('validation_results') or section.get('metadata', {}).get('validation_results', {})
        
        # Section-level validation summary at the top
        if validation_data:
            # Compact section metrics in a single row
            col1, col2, col3, col4 = st.columns(4)
            
            quality_score = validation_data.get('overall_quality_score') or validation_data.get('confidence_score', 0.0)
            total_issues = validation_data.get('total_issues')
            if total_issues is None:
                issues = validation_data.get('issues', [])
                total_issues = len(issues)
            
            with col1:
                if quality_score >= 0.8:
                    st.success(f"üìä Quality: {quality_score:.0%}")
                elif quality_score >= 0.6:
                    st.warning(f"üìä Quality: {quality_score:.0%}")
                else:
                    st.error(f"üìä Quality: {quality_score:.0%}")
            
            with col2:
                st.info(f"üìã Issues: {total_issues}")
            
            with col3:
                issues_by_severity = validation_data.get('issues_by_severity', {})
                high_issues = issues_by_severity.get('high', 0)
                if high_issues > 0:
                    st.error(f"üî¥ High: {high_issues}")
                else:
                    st.success("üü¢ No High Issues")
            
            with col4:
                assessment = validation_data.get('overall_assessment', 'Analysis completed')
                status_text = "‚úÖ Passed" if quality_score >= 0.8 else "‚ö†Ô∏è Review Required" if quality_score >= 0.6 else "‚ùå Needs Revision"
                st.write(f"**{status_text}**")
        
        st.markdown("---")
        
        # Display content paragraph by paragraph with side-by-side validation
        content = section.get('content', '')
        if content:
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            
            for p_idx, paragraph in enumerate(paragraphs):
                # Create a container for each paragraph with side-by-side layout
                st.markdown(f"### üìÑ Paragraph {p_idx + 1}")
                
                # Split into two columns for content and validation
                content_col, validation_col = st.columns([1.3, 0.7])
                
                with content_col:
                    # Display paragraph content in a text box
                    st.text_area(
                        label="Generated Content",
                        value=paragraph,
                        height=150,
                        disabled=True,
                        key=f"content_{i}_{p_idx}",
                        label_visibility="visible"
                    )
                
                with validation_col:
                    # Get validation for this paragraph
                    validation_status = get_paragraph_validation(validation_data, paragraph)
                    status = validation_status.get('status', 'unknown')
                    conclusion = validation_status.get('conclusion', 'Unknown status')
                    issues = validation_status.get('issues', [])
                    
                    # Status indicator at the top
                    if status == 'passed':
                        st.success(f"‚úÖ Validation Passed")
                        st.write(f"_{conclusion}_")
                    elif status == 'partially_passed':
                        st.warning(f"‚ö†Ô∏è Partial Pass")
                        st.write(f"_{conclusion}_")
                    elif status == 'not_passed':
                        st.error(f"‚ùå Failed Validation")
                        st.write(f"_{conclusion}_")
                    else:
                        st.info(f"‚ÑπÔ∏è Unknown Status")
                        st.write(f"_{conclusion}_")
                    
                    # Display issues in a compact format
                    if issues:
                        st.markdown("**Validation Issues:**")
                        for issue in issues:
                            severity = issue.get('severity', 'medium')
                            severity_icon = {
                                'high': 'üî¥',
                                'medium': 'üü°',
                                'low': 'üü¢',
                                'info': 'üîµ'
                            }.get(severity, '‚ö™')
                            
                            issue_type = issue.get('issue_type', 'Issue')
                            description = issue.get('description', 'No description')
                            
                            # Create expandable details for each issue
                            with st.expander(f"{severity_icon} {issue_type.title()}", expanded=(severity == 'high')):
                                st.write(f"**Severity:** {severity.title()}")
                                st.write(f"**Description:** {description}")
                                
                                if issue.get('text_span'):
                                    st.write(f"**Affected Text:** _{issue.get('text_span')}_")
                                
                                if issue.get('suggested_fix'):
                                    st.write(f"**üí° Suggested Fix:** {issue.get('suggested_fix')}")
                                
                                confidence = issue.get('confidence_score', 0.0)
                                st.write(f"**AI Confidence:** {confidence:.1%}")
                    else:
                        st.success("**‚úÖ No Issues Found**")
                        st.write("This paragraph passed all validation checks.")
                
                # Add spacing between paragraphs
                if p_idx < len(paragraphs) - 1:
                    st.markdown("---")
        else:
            st.warning("No content available for this section")
        
        # Final section separator
        st.markdown("---")
    
    # Download options - ALWAYS VISIBLE
    st.subheader("üíæ Download Report")

    # Persistent status indicator
    if hasattr(st.session_state, 'word_preparation_state'):
        if st.session_state.word_preparation_state == 'preparing':
            st.info("üîÑ **Status:** Word document is being prepared... Please wait.")
        elif st.session_state.word_preparation_state == 'prepared':
            st.success("‚úÖ **Status:** Word document is ready for download!")
        elif st.session_state.word_preparation_state == 'error':
            st.error("‚ùå **Status:** Document preparation encountered errors.")

    # Download section stays visible always - create 3 columns for all download options
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # Word Document Download Section - Fixed approach without fragment issues
        word_download_section(report)
    
    with col2:
        # Text Download Section using st.fragment - prevents full page reload
        text_download_fragment(report)

    with col3:
        # JSON Download Section using st.fragment - prevents full page reload
        json_download_fragment(report)

if __name__ == "__main__":
    main()