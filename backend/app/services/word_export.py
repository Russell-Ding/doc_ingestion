import asyncio
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import structlog
from datetime import datetime

from app.core.config import settings
# from app.services.validation_system import ValidationIssue, ValidationSeverity
# Import removed to avoid dependency issues - validation types are used as dicts

logger = structlog.get_logger(__name__)


class WordReportGenerator:
    """Generates Word documents with embedded validation comments"""
    
    def __init__(self):
        self.template_path = settings.WORD_TEMPLATE_PATH
        self.export_directory = Path(settings.EXPORT_DIRECTORY)
        # Ensure export directory exists with full permissions
        self.export_directory.mkdir(parents=True, exist_ok=True)
        logger.info(f"Export directory initialized: {self.export_directory.absolute()}")
    
    async def generate_report(
        self,
        report_data: Dict[str, Any],
        segments: List[Dict[str, Any]],
        validation_results: List[Dict[str, Any]],
        include_validation_comments: bool = True
    ) -> Dict[str, str]:
        """Generate a complete Word report with validation comments"""
        
        logger.info(
            "Starting Word report generation",
            report_id=report_data.get("id"),
            segment_count=len(segments),
            include_comments=include_validation_comments
        )
        
        try:
            # Create new document or load template
            doc = self._create_document()
            
            # Set up document styles
            self._setup_document_styles(doc)
            
            # Add document header
            self._add_document_header(doc, report_data)
            
            # Add executive summary
            self._add_executive_summary(doc, report_data, segments)
            
            # Add report segments
            for i, segment in enumerate(segments):
                self._add_segment_content(
                    doc, 
                    segment, 
                    validation_results[i] if i < len(validation_results) else None,
                    include_validation_comments
                )
            
            # Add appendices
            self._add_appendices(doc, report_data, segments)
            
            # Save document
            filename = self._generate_filename(report_data)
            file_path = self.export_directory / filename
            
            doc.save(str(file_path))
            
            logger.info(
                "Word report generated successfully",
                filename=filename,
                file_size=file_path.stat().st_size
            )
            
            return {
                "filename": filename,
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size,
                "download_url": f"/api/v1/reports/download/{filename}"
            }
            
        except Exception as e:
            logger.error("Word report generation failed", error=str(e))
            raise
    
    def _create_document(self) -> Document:
        """Create a new Word document or load from template"""
        if self.template_path and Path(self.template_path).exists():
            return Document(self.template_path)
        else:
            return Document()
    
    def _setup_document_styles(self, doc: Document):
        """Set up custom styles for the document"""
        styles = doc.styles
        
        # Title style
        if 'Credit Report Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Credit Report Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(18)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0, 51, 102)
            
            title_paragraph = title_style.paragraph_format
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_paragraph.space_after = Pt(12)
        
        # Section heading style
        if 'Section Heading' not in [s.name for s in styles]:
            heading_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Arial'
            heading_font.size = Pt(14)
            heading_font.bold = True
            heading_font.color.rgb = RGBColor(0, 51, 102)
            
            heading_paragraph = heading_style.paragraph_format
            heading_paragraph.space_before = Pt(18)
            heading_paragraph.space_after = Pt(6)
        
        # Validation comment style
        if 'Validation Issue' not in [s.name for s in styles]:
            validation_style = styles.add_style('Validation Issue', WD_STYLE_TYPE.CHARACTER)
            validation_font = validation_style.font
            validation_font.name = 'Arial'
            validation_font.size = Pt(9)
            validation_font.italic = True
    
    def _add_document_header(self, doc: Document, report_data: Dict[str, Any]):
        """Add document header with title and metadata"""
        
        # Main title
        title_paragraph = doc.add_paragraph()
        title_paragraph.style = 'Credit Report Title'
        title_run = title_paragraph.add_run(report_data.get('title', 'Credit Review Report'))
        
        # Subtitle with date
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.italic = True
        
        # Add page break
        doc.add_page_break()
    
    def _add_executive_summary(
        self, 
        doc: Document, 
        report_data: Dict[str, Any], 
        segments: List[Dict[str, Any]]
    ):
        """Add executive summary section"""
        
        # Executive Summary heading
        heading = doc.add_paragraph()
        heading.style = 'Section Heading'
        heading.add_run('Executive Summary')
        
        # Summary content
        summary_paragraph = doc.add_paragraph()
        summary_text = self._generate_executive_summary(report_data, segments)
        summary_paragraph.add_run(summary_text)
        
        # Key findings table
        self._add_key_findings_table(doc, segments)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_segment_content(
        self,
        doc: Document,
        segment: Dict[str, Any],
        validation_result: Optional[Dict[str, Any]],
        include_validation_comments: bool
    ):
        """Add individual segment content with validation comments"""
        
        # Section heading
        heading = doc.add_paragraph()
        heading.style = 'Section Heading'
        heading.add_run(segment.get('name', 'Untitled Section'))
        
        # Segment content
        content = segment.get('generated_content', '')
        if not content:
            placeholder = doc.add_paragraph()
            placeholder.add_run('[Content not generated for this section]')
            placeholder.runs[0].font.italic = True
            placeholder.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            return
        
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                paragraph = doc.add_paragraph()
                
                if include_validation_comments and validation_result:
                    self._add_content_with_comments(
                        paragraph, 
                        paragraph_text, 
                        validation_result.get('word_comments', [])
                    )
                else:
                    paragraph.add_run(paragraph_text)
        
        # Add validation summary if available
        if include_validation_comments and validation_result:
            self._add_validation_summary_box(doc, validation_result)
        
        doc.add_paragraph()  # Add spacing
    
    def _add_content_with_comments(
        self,
        paragraph,
        content: str,
        word_comments: List[Dict[str, Any]]
    ):
        """Add content with embedded validation comments"""
        
        if not word_comments:
            paragraph.add_run(content)
            return
        
        # Find comments that apply to this paragraph content
        relevant_comments = []
        for comment in word_comments:
            # Try to get text span from different possible fields
            text_span = comment.get('text_span', comment.get('text', ''))
            
            # If we have a text span and it's in this paragraph content
            if text_span and text_span.strip() in content:
                # Find the position of this text span in the current paragraph
                start_pos = content.find(text_span.strip())
                if start_pos != -1:
                    relevant_comment = comment.copy()
                    relevant_comment['start'] = start_pos
                    relevant_comment['end'] = start_pos + len(text_span.strip())
                    relevant_comments.append(relevant_comment)
            else:
                # If no specific text span, add comment to the beginning of paragraph
                relevant_comment = comment.copy()
                relevant_comment['start'] = 0
                relevant_comment['end'] = min(50, len(content))  # First 50 chars
                relevant_comments.append(relevant_comment)
        
        if not relevant_comments:
            paragraph.add_run(content)
            return
        
        # Sort comments by position
        sorted_comments = sorted(relevant_comments, key=lambda x: x.get('start', 0))
        
        current_pos = 0
        
        for comment in sorted_comments:
            start_pos = comment.get('start', 0)
            end_pos = comment.get('end', start_pos + 1)
            
            # Ensure positions are within content bounds
            start_pos = max(0, min(start_pos, len(content)))
            end_pos = max(start_pos, min(end_pos, len(content)))
            
            # Add text before the comment
            if current_pos < start_pos:
                paragraph.add_run(content[current_pos:start_pos])
            
            # Add highlighted text with comment
            if start_pos < end_pos:
                highlighted_text = content[start_pos:end_pos]
                comment_run = paragraph.add_run(highlighted_text)
                
                # Apply highlighting based on severity
                severity = comment.get('severity', 'medium')
                if severity == 'high':
                    comment_run.font.highlight_color = WD_COLOR_INDEX.RED
                elif severity == 'medium':
                    comment_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                elif severity == 'low':
                    comment_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                else:
                    comment_run.font.highlight_color = WD_COLOR_INDEX.BLUE
                
                # Add comment as footnote annotation
                comment_text = comment.get('text', '')
                comment_type = comment.get('type', 'validation')

                # Ensure comment text is not empty
                if comment_text:
                    footnote_run = paragraph.add_run(f" [AI Validation - {comment_type.title()}: {comment_text}]")
                    footnote_run.font.size = Pt(8)
                    footnote_run.font.italic = True

                    # Color code the comment based on severity
                    if severity == 'high':
                        footnote_run.font.color.rgb = RGBColor(204, 0, 0)  # Dark red
                    elif severity == 'medium':
                        footnote_run.font.color.rgb = RGBColor(255, 102, 0)  # Orange
                    else:
                        footnote_run.font.color.rgb = RGBColor(0, 102, 204)  # Blue
            
            current_pos = end_pos
        
        # Add remaining text
        if current_pos < len(content):
            paragraph.add_run(content[current_pos:])
        
        # Add paragraph-level validation summary if there are comments
        if relevant_comments:
            validation_summary_run = paragraph.add_run(
                f"\n[Paragraph Validation: {len(relevant_comments)} issue(s) identified]"
            )
            validation_summary_run.font.size = Pt(9)
            validation_summary_run.font.italic = True
            validation_summary_run.font.color.rgb = RGBColor(102, 102, 102)
    
    def _add_validation_summary_box(
        self,
        doc: Document,
        validation_result: Dict[str, Any]
    ):
        """Add a summary box showing validation results"""
        
        summary = validation_result.get('summary', {})
        total_issues = validation_result.get('total_issues', 0)
        
        if total_issues == 0:
            return
        
        # Add a bordered text box (simplified approach using table)
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        cell = table.cell(0, 0)
        cell_paragraph = cell.paragraphs[0]
        
        # Title
        title_run = cell_paragraph.add_run('Validation Summary')
        title_run.font.bold = True
        title_run.font.size = Pt(11)
        
        cell_paragraph.add_run(f"\n{total_issues} issues identified")
        
        # Quality score
        quality_score = validation_result.get('overall_quality_score', 0.0)
        quality_rating = summary.get('quality_rating', 'Unknown')
        cell_paragraph.add_run(f"\nQuality Score: {quality_score:.1%} ({quality_rating})")
        
        # Issue breakdown
        issues_by_severity = validation_result.get('issues_by_severity', {})
        if any(issues_by_severity.values()):
            cell_paragraph.add_run("\nIssue Breakdown:")
            for severity, count in issues_by_severity.items():
                if count > 0:
                    cell_paragraph.add_run(f"\n• {severity.title()}: {count}")
        
        # Set cell background color
        cell._element.get_or_add_tcPr().append(
            self._create_cell_color_element('F0F0F0')
        )
        
        doc.add_paragraph()  # Add spacing
    
    def _add_key_findings_table(self, doc: Document, segments: List[Dict[str, Any]]):
        """Add a table summarizing key findings from each segment"""
        
        if not segments:
            return
        
        # Table heading
        table_heading = doc.add_paragraph()
        table_run = table_heading.add_run('Key Findings Summary')
        table_run.font.bold = True
        table_run.font.size = Pt(12)
        
        # Create table
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Section'
        header_cells[1].text = 'Status'
        header_cells[2].text = 'Key Points'
        
        # Format header
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add data rows
        for segment in segments:
            row_cells = table.add_row().cells
            
            row_cells[0].text = segment.get('name', 'Untitled')
            
            # Status based on content_status
            status = segment.get('content_status', 'pending')
            status_display = {
                'completed': '✓ Complete',
                'generating': '⏳ Generating',
                'error': '❌ Error',
                'pending': '⏸ Pending'
            }.get(status, status.title())
            
            row_cells[1].text = status_display
            
            # Extract key points from content (simplified)
            content = segment.get('generated_content', '')
            key_points = self._extract_key_points(content)
            row_cells[2].text = key_points
        
        doc.add_paragraph()  # Add spacing
    
    def _add_appendices(
        self, 
        doc: Document, 
        report_data: Dict[str, Any], 
        segments: List[Dict[str, Any]]
    ):
        """Add appendices with source documents and methodology"""
        
        doc.add_page_break()
        
        # Appendices heading
        heading = doc.add_paragraph()
        heading.style = 'Section Heading'
        heading.add_run('Appendices')
        
        # Appendix A: Source Documents
        self._add_source_documents_appendix(doc, segments)
        
        # Appendix B: Methodology
        self._add_methodology_appendix(doc)
        
        # Appendix C: Validation Details (if validation was performed)
        if any(segment.get('validation_result') for segment in segments):
            self._add_validation_details_appendix(doc, segments)
    
    def _add_source_documents_appendix(self, doc: Document, segments: List[Dict[str, Any]]):
        """Add appendix listing source documents used"""
        
        subheading = doc.add_paragraph()
        subheading_run = subheading.add_run('Appendix A: Source Documents')
        subheading_run.font.bold = True
        subheading_run.font.size = Pt(12)
        
        # Collect unique documents from all segments
        all_documents = set()
        for segment in segments:
            references = segment.get('references', [])
            for ref in references:
                doc_id = ref.get('document_id')
                if doc_id:
                    all_documents.add(doc_id)
        
        if all_documents:
            doc_list = doc.add_paragraph()
            doc_list.add_run("The following documents were referenced in this analysis:")
            
            for i, doc_id in enumerate(sorted(all_documents), 1):
                doc_item = doc.add_paragraph()
                doc_item.add_run(f"{i}. Document ID: {doc_id}")
        else:
            no_docs = doc.add_paragraph()
            no_docs.add_run("No specific source documents were referenced.")
        
        doc.add_paragraph()  # Add spacing
    
    def _add_methodology_appendix(self, doc: Document):
        """Add methodology appendix"""
        
        subheading = doc.add_paragraph()
        subheading_run = subheading.add_run('Appendix B: Methodology')
        subheading_run.font.bold = True
        subheading_run.font.size = Pt(12)
        
        methodology_text = """
This credit review report was generated using an AI-powered document analysis system. The methodology includes:

1. Document Ingestion: All provided documents were processed and indexed for content retrieval.

2. Semantic Analysis: Advanced natural language processing techniques were used to understand document contents and extract relevant information.

3. Content Generation: AI models analyzed the source documents and generated report content based on the specified requirements for each section.

4. Validation: Generated content was validated against source documents to ensure accuracy and completeness.

5. Quality Assurance: Multiple validation checks were performed to maintain report quality and compliance standards.

The system uses state-of-the-art AI models while maintaining human oversight and validation throughout the process.
        """
        
        methodology_paragraph = doc.add_paragraph()
        methodology_paragraph.add_run(methodology_text.strip())
        
        doc.add_paragraph()  # Add spacing
    
    def _add_validation_details_appendix(self, doc: Document, segments: List[Dict[str, Any]]):
        """Add detailed validation results appendix"""
        
        subheading = doc.add_paragraph()
        subheading_run = subheading.add_run('Appendix C: Validation Details')
        subheading_run.font.bold = True
        subheading_run.font.size = Pt(12)
        
        for segment in segments:
            validation_result = segment.get('validation_result')
            if not validation_result:
                continue
            
            # Section name
            section_paragraph = doc.add_paragraph()
            section_run = section_paragraph.add_run(f"Section: {segment.get('name', 'Untitled')}")
            section_run.font.bold = True
            
            # Validation summary
            total_issues = validation_result.get('total_issues', 0)
            quality_score = validation_result.get('overall_quality_score', 0.0)
            
            summary_paragraph = doc.add_paragraph()
            summary_paragraph.add_run(f"Total Issues: {total_issues}, Quality Score: {quality_score:.1%}")
            
            # Issue details
            if total_issues > 0:
                issues_by_type = validation_result.get('issues_by_type', {})
                for issue_type, count in issues_by_type.items():
                    if count > 0:
                        issue_paragraph = doc.add_paragraph()
                        issue_paragraph.add_run(f"• {issue_type.title()}: {count}")
            
            doc.add_paragraph()  # Add spacing between sections
    
    def _generate_executive_summary(
        self, 
        report_data: Dict[str, Any], 
        segments: List[Dict[str, Any]]
    ) -> str:
        """Generate executive summary text"""
        
        completed_segments = [s for s in segments if s.get('content_status') == 'completed']
        total_segments = len(segments)
        
        summary = f"""
This credit review report provides a comprehensive analysis based on the submitted documentation. 
The analysis covers {total_segments} key areas, with {len(completed_segments)} sections completed.

Report generated using AI-powered document analysis to ensure comprehensive coverage of all 
relevant factors affecting the credit decision. Each section has been validated against 
source documents to maintain accuracy and reliability.

Key areas of analysis include financial performance, risk assessment, industry conditions, 
and other factors material to the credit evaluation process.
        """
        
        return summary.strip()
    
    def _extract_key_points(self, content: str) -> str:
        """Extract key points from segment content"""
        if not content:
            return "No content available"
        
        # Simple extraction of first sentence or key phrases
        sentences = content.split('. ')
        if sentences:
            first_sentence = sentences[0].strip()
            return first_sentence[:100] + "..." if len(first_sentence) > 100 else first_sentence
        
        return content[:100] + "..." if len(content) > 100 else content
    
    def _generate_filename(self, report_data: Dict[str, Any]) -> str:
        """Generate appropriate filename for the report"""
        
        title = report_data.get('title', 'Credit_Review_Report')
        # Clean title for filename
        clean_title = ''.join(c for c in title if c.isalnum() or c in (' ', '-', '_'))
        clean_title = clean_title.replace(' ', '_')
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        return f"{clean_title}_{timestamp}.docx"
    
    def _create_cell_color_element(self, color: str):
        """Create XML element for cell background color"""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), color)
        return shading


# Global Word export service instance
word_export_service = WordReportGenerator()